from docx import Document
import re

path = '/Users/piyushsingh/Downloads/HECA PROJECT /HECA_Research_Paper_Singh_Chatterjee.docx'
doc = Document(path)

EM = '\u2014'   # em dash  —
EN = '\u2013'   # en dash  –

def replace_all_dashes(text):
    if EM not in text and EN not in text:
        return text

    # 1. Numeric ranges: 40–70  or  40—70  →  40-70  (hyphen, no spaces)
    text = re.sub(r'(\d)\s*[' + EM + EN + r']\s*(\d)', r'\1-\2', text)

    # 2. "Layer N – Title"  →  "Layer N: Title"
    text = re.sub(r'(Layer\s*\d+)\s*[' + EM + EN + r']\s*', r'\1: ', text)

    # 3. Double-dash parenthetical  "– ... –"  →  ", ... ,"
    def paren(m):
        return ', ' + m.group(1).strip() + ','
    for dash in [EN, EM]:
        text = re.sub(
            re.escape(dash) + r'([^' + EM + EN + r']{3,80}?)' + re.escape(dash),
            paren, text
        )

    # 4. After linking verbs  →  colon
    text = re.sub(
        r'\b(is|was|are|were|has|have|had|means|reveals|shows|remains|does|did)'
        r'\s*[' + EM + EN + r']\s*',
        r'\1: ', text
    )

    # 5.  "– not"  →  ", not"
    text = re.sub(r'\s*[' + EM + EN + r']\s*(not\b)', r', \1', text)

    # 6. Everything else: spaced dash → comma (most readable for academic prose)
    #    "word – phrase" → "word, phrase"
    text = re.sub(r'\s*[' + EM + EN + r']\s*', ', ', text)

    # Clean up any double commas or comma after opening bracket
    text = re.sub(r',\s*,', ',', text)
    text = re.sub(r'\(\s*,\s*', '(', text)
    text = re.sub(r',\s*\)', ')', text)
    text = re.sub(r'  +', ' ', text)

    return text

def fix_runs(para):
    for run in para.runs:
        if EM in run.text or EN in run.text:
            run.text = replace_all_dashes(run.text)

# Paragraphs
for para in doc.paragraphs:
    fix_runs(para)

# Tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                fix_runs(para)

doc.save(path)

# Final audit
em_left = sum(p.text.count(EM) for p in doc.paragraphs)
en_left  = sum(p.text.count(EN) for p in doc.paragraphs)
print(f'Em dashes remaining: {em_left}')
print(f'En dashes remaining: {en_left}')
print('Done.')
