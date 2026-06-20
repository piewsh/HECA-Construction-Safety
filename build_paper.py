"""
HECA Research Paper — Professional Word Document Builder
Piyush Singh & Abhishek Chatterjee | IIT Madras
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, copy

# ─────────────────────────────────────────────────────────────────────────────
# LOW-LEVEL HELPERS
# ─────────────────────────────────────────────────────────────────────────────

NAVY   = RGBColor(0x1B, 0x26, 0x4E)   # deep navy for headings
BLACK  = RGBColor(0x0D, 0x0D, 0x0D)   # near-black for body
GREY   = RGBColor(0x66, 0x66, 0x66)   # grey for captions / placeholders
TABLE_HDR = 'D6DCF0'                   # light blue-grey for table header
TABLE_ROW = 'F5F6FB'                   # very light for alternate rows

def set_run_font(run, name='Times New Roman', size=12, bold=False,
                 italic=False, color=BLACK):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color

def para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=6,
             line_spacing=WD_LINE_SPACING.DOUBLE,
             left_indent=0, first_indent=0):
    pf = p.paragraph_format
    pf.alignment        = align
    pf.space_before     = Pt(space_before)
    pf.space_after      = Pt(space_after)
    pf.line_spacing_rule = line_spacing
    if left_indent:
        pf.left_indent  = Inches(left_indent)
    if first_indent:
        pf.first_line_indent = Inches(first_indent)

def add_border(para, side='bottom', size=8, color='1B264E', space=4):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement(f'w:{side}')
    el.set(qn('w:val'),   'single')
    el.set(qn('w:sz'),    str(size))
    el.set(qn('w:space'), str(space))
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:color'), 'AAAACC')
        tblBdr.append(el)
    tblPr.append(tblBdr)

# ─────────────────────────────────────────────────────────────────────────────
# INLINE MARKUP RENDERER  (**bold**, *italic*, plain)
# ─────────────────────────────────────────────────────────────────────────────

EM = '\u2014'   # em dash

def clean(text):
    """Normalise punctuation."""
    text = text.replace('—', EM).replace('--', EM)
    text = text.replace(' - ', f'\u202f{EM}\u202f')   # thin-spaced em dash
    text = text.replace('...', '\u2026')
    text = text.replace("'", '\u2019').replace('"', '\u201c').replace('"', '\u201d')
    return text

def render_inline(para, text, base_size=12, base_bold=False, base_italic=False,
                  base_color=BLACK):
    """Parse **bold**, *italic*, `code` and add runs to para."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))', re.DOTALL)
    for m in pattern.finditer(text):
        if m.group(2):          # **bold**
            r = para.add_run(clean(m.group(2)))
            set_run_font(r, size=base_size, bold=True,
                         italic=base_italic, color=base_color)
        elif m.group(3):        # *italic*
            r = para.add_run(clean(m.group(3)))
            set_run_font(r, size=base_size, bold=base_bold,
                         italic=True, color=base_color)
        elif m.group(4):        # `code`
            r = para.add_run(m.group(4))
            r.font.name = 'Courier New'
            r.font.size = Pt(base_size - 1)
        elif m.group(5):        # plain
            r = para.add_run(clean(m.group(5)))
            set_run_font(r, size=base_size, bold=base_bold,
                         italic=base_italic, color=base_color)

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Page: US Letter, 1.25" L/R, 1.0" T/B
for s in doc.sections:
    s.page_width    = Inches(8.5)
    s.page_height   = Inches(11)
    s.left_margin   = Inches(1.25)
    s.right_margin  = Inches(1.25)
    s.top_margin    = Inches(1.0)
    s.bottom_margin = Inches(1.0)

# ── Global style overrides ────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

for lvl, (sz, bold, italic) in enumerate(
        [(16,True,False),(13,True,False),(12,True,True)], start=1):
    h = doc.styles[f'Heading {lvl}']
    h.font.name   = 'Times New Roman'
    h.font.size   = Pt(sz)
    h.font.bold   = bold
    h.font.italic = italic
    h.font.color.rgb = NAVY
    h.paragraph_format.space_before = Pt([24,16,12][lvl-1])
    h.paragraph_format.space_after  = Pt(6)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h.paragraph_format.keep_with_next    = True

# ─────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────

def blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        para_fmt(p, space_before=0, space_after=0,
                 line_spacing=WD_LINE_SPACING.SINGLE)

blank(4)

# TITLE
p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0,
         space_after=20, line_spacing=WD_LINE_SPACING.SINGLE)
r = p.add_run(
    'Measuring What Kills:\n'
    'Energy-Based Safety Assessment of Indian PPVC Construction\n'
    'Using the HECA Framework'
)
set_run_font(r, size=20, bold=True, color=NAVY)

# thin rule under title
p_rule = doc.add_paragraph()
para_fmt(p_rule, space_before=0, space_after=16,
         line_spacing=WD_LINE_SPACING.SINGLE)
add_border(p_rule, 'bottom', size=12, color='1B264E')

# Authors
p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0,
         space_after=6, line_spacing=WD_LINE_SPACING.SINGLE)
r = p.add_run('Piyush Singh\u00b9  \u00b7  Abhishek Chatterjee\u00b9')
set_run_font(r, size=13, bold=False, color=BLACK)

# Affiliation
p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0,
         space_after=4, line_spacing=WD_LINE_SPACING.SINGLE)
r = p.add_run(
    '\u00b9 Department of Civil Engineering\n'
    'Indian Institute of Technology Madras, Chennai 600036, India'
)
set_run_font(r, size=10, italic=True, color=GREY)

# Correspondence
p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0,
         space_after=8, line_spacing=WD_LINE_SPACING.SINGLE)
r = p.add_run('Correspondence: piyushsingh@iitm.ac.in')
set_run_font(r, size=10, color=GREY)

blank(2)

# Journal note
p = doc.add_paragraph()
para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0,
         space_after=0, line_spacing=WD_LINE_SPACING.SINGLE)
r = p.add_run('Prepared for submission to Safety Science / Automation in Construction')
set_run_font(r, size=10, italic=True, color=GREY)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# ABSTRACT
# ─────────────────────────────────────────────────────────────────────────────

h = doc.add_heading('Abstract', level=1)

ABSTRACT = (
    "India\u2019s construction sector accounts for approximately one in every four occupational "
    "fatalities in the country, yet no energy-based safety baseline exists for its rapidly expanding "
    "Prefabricated Prefinished Volumetric Construction (PPVC) sector. This study deploys the "
    "High-Energy Control Assessment (HECA) framework\u2014an energy-based safety monitoring tool "
    "developed in North America and never previously applied in India\u2014to an active PPVC project "
    "at the Indian Institute of Technology Madras. Through a non-invasive photographic audit of 320 "
    "curated images, two trained assessors identified 1,583 hazard observations and classified "
    "1,254 (79.2%) as high-energy with serious injury and fatality (SIF) potential. The aggregate "
    "HECA score was 1.83%, against the 40\u201370% benchmark range reported for North American utility "
    "projects. In 297 of 320 images (92.8%), every visible high-energy hazard was completely "
    "uncontrolled. The PPVC Module Erection phase recorded 0.00% across 189 high-energy hazard "
    "observations; five of seven energy types recorded exactly zero. Beyond the field analysis, the "
    "study delivers two technical contributions: a six-layer Image Anonymization Pipeline achieving "
    "compliance with India\u2019s Digital Personal Data Protection (DPDP) Act 2023, and a "
    "zero-dependency HECA Analytics Dashboard enabling site-level data exploration without "
    "institutional IT infrastructure. Together, these constitute the first complete, replicable "
    "HECA research system for Indian construction."
)

p = doc.add_paragraph()
para_fmt(p, space_before=0, space_after=8,
         line_spacing=WD_LINE_SPACING.SINGLE)
p.paragraph_format.left_indent  = Inches(0.35)
p.paragraph_format.right_indent = Inches(0.35)
r = p.add_run(ABSTRACT)
set_run_font(r, size=11, color=BLACK)

# Keywords
p = doc.add_paragraph()
para_fmt(p, space_before=4, space_after=20,
         line_spacing=WD_LINE_SPACING.SINGLE)
p.paragraph_format.left_indent  = Inches(0.35)
p.paragraph_format.right_indent = Inches(0.35)
r = p.add_run('Keywords: ')
set_run_font(r, size=11, bold=True)
r = p.add_run(
    'HECA \u00b7 High-Energy Control Assessment \u00b7 PPVC \u00b7 '
    'Modular Construction \u00b7 Direct Controls \u00b7 SIF Prevention \u00b7 '
    'Photographic Audit \u00b7 Image Anonymization \u00b7 DPDP Act 2023 \u00b7 '
    'India \u00b7 Construction Safety'
)
set_run_font(r, size=11, italic=True)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# BODY CONTENT  —  read and render the markdown
# ─────────────────────────────────────────────────────────────────────────────

with open(
    '/Users/piyushsingh/.gemini/antigravity/brain/'
    '50562d44-3450-4c11-945c-c1fe6417aacf/RESEARCH_PAPER_HUMAN_DRAFT.md',
    encoding='utf-8'
) as f:
    raw = f.read()

# ── Pre-processing: strip title block, abstract block, and trailing note ─────

# Drop everything up to and including "## 1. INTRODUCTION"
raw = re.sub(r'^.*?(## 1\. INTRODUCTION)', r'\1', raw, flags=re.DOTALL)
# Drop the footer note
raw = re.sub(r'\*\[Complete paper.*', '', raw, flags=re.DOTALL)
# Drop the trailing REFERENCES header (we'll rebuild it clean)
# We'll keep refs content but re-render it

# ── Split into blocks ─────────────────────────────────────────────────────────

blocks = []
current = []

def flush(buf):
    txt = '\n'.join(buf).strip()
    if txt:
        blocks.append(txt)

for line in raw.splitlines():
    s = line.strip()
    if s.startswith('#'):
        flush(current); current = [line]
    elif not s:
        flush(current); current = []
    else:
        current.append(line)
flush(current)

# ── Table detector ────────────────────────────────────────────────────────────

def is_table_block(blk):
    lines = blk.strip().splitlines()
    return lines and lines[0].strip().startswith('|')

def parse_table(blk):
    rows_raw = [l for l in blk.strip().splitlines() if l.strip().startswith('|')]
    rows = []
    for r in rows_raw:
        cells = [c.strip() for c in r.strip('|').split('|')]
        if any(re.match(r'^[-: ]+$', c) for c in cells):
            continue   # separator row
        rows.append(cells)
    return rows

def add_docx_table(rows):
    if not rows:
        return
    ncols  = max(len(r) for r in rows)
    table  = doc.add_table(rows=len(rows), cols=ncols)
    set_table_borders(table)
    table.autofit = True

    for ri, row_data in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            txt  = row_data[ci] if ci < len(row_data) else ''
            # strip bold/italic markers for cell content
            plain = re.sub(r'\*+(.+?)\*+', r'\1', txt)
            plain = clean(plain)
            cell.text = ''
            cp = cell.paragraphs[0]
            para_fmt(cp, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=3, space_after=3,
                     line_spacing=WD_LINE_SPACING.SINGLE)
            r = cp.add_run(plain)
            set_run_font(r, size=10, bold=(ri == 0), color=BLACK)
            shade_cell(cell, TABLE_HDR if ri == 0
                       else (TABLE_ROW if ri % 2 == 0 else 'FFFFFF'))

    # space after table
    p = doc.add_paragraph()
    para_fmt(p, space_before=0, space_after=8,
             line_spacing=WD_LINE_SPACING.SINGLE)

# ── Block renderer ────────────────────────────────────────────────────────────

ref_mode = False

for blk in blocks:
    s = blk.strip()
    if not s:
        continue

    # ── Horizontal rules ──────────────────────────────────────────────────────
    if s == '---':
        p = doc.add_paragraph()
        para_fmt(p, space_before=4, space_after=4,
                 line_spacing=WD_LINE_SPACING.SINGLE)
        add_border(p, 'bottom', size=4, color='CCCCDD')
        continue

    # ── REFERENCES section ────────────────────────────────────────────────────
    if s.startswith('## REFERENCES'):
        ref_mode = True
        doc.add_page_break()
        doc.add_heading('References', level=1)
        continue

    if ref_mode:
        # Each reference is [N] Author...
        m = re.match(r'^\[(\d+)\]\s+(.*)', s, re.DOTALL)
        if m:
            num, body = m.group(1), m.group(2)
            p = doc.add_paragraph()
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     space_before=0, space_after=5,
                     line_spacing=WD_LINE_SPACING.SINGLE)
            p.paragraph_format.left_indent       = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.4)
            # number
            r = p.add_run(f'[{num}] ')
            set_run_font(r, size=11, bold=True, color=NAVY)
            # body — handle *italic journal names*
            parts = re.split(r'(\*[^*]+\*)', body)
            for part in parts:
                if part.startswith('*') and part.endswith('*'):
                    r = p.add_run(clean(part[1:-1]))
                    set_run_font(r, size=11, italic=True)
                else:
                    r = p.add_run(clean(part))
                    set_run_font(r, size=11)
        continue

    # ── Headings ──────────────────────────────────────────────────────────────
    if s.startswith('#### '):
        doc.add_heading(clean(s[5:]), level=3)
        continue
    if s.startswith('### '):
        doc.add_heading(clean(s[4:]), level=2)
        continue
    if s.startswith('## '):
        text = clean(s[3:])
        h = doc.add_heading(text, level=1)
        continue
    if s.startswith('# ') and not s.startswith('## '):
        # top-level title — skip (already on title page)
        continue

    # ── Equation line (centred, bold) ─────────────────────────────────────────
    if 'HECA Score (%)' in s and '=' in s:
        eq = re.sub(r'\*+', '', s).strip()
        p  = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=10, space_after=10,
                 line_spacing=WD_LINE_SPACING.SINGLE)
        r = p.add_run(clean(eq))
        set_run_font(r, size=12, bold=True)
        continue

    # ── Figure placeholder ────────────────────────────────────────────────────
    if re.match(r'^\*?\[Figure', s) or re.match(r'^\*?\[Fig\.', s):
        txt = re.sub(r'[\*\[\]\n]', '', s).strip()
        p   = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=6, space_after=6,
                 line_spacing=WD_LINE_SPACING.SINGLE)
        r = p.add_run(f'〔 {txt} 〕')
        set_run_font(r, size=10, italic=True, color=GREY)
        continue

    # ── Table ─────────────────────────────────────────────────────────────────
    if is_table_block(s):
        rows = parse_table(s)
        add_docx_table(rows)
        continue

    # ── Bullet list ───────────────────────────────────────────────────────────
    if re.match(r'^[-*•] ', s):
        for line in s.splitlines():
            ls = line.strip()
            if not ls or not re.match(r'^[-*•] ', ls):
                continue
            item = ls[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     space_before=0, space_after=3,
                     line_spacing=WD_LINE_SPACING.DOUBLE)
            p.paragraph_format.left_indent = Inches(0.35)
            render_inline(p, item, base_size=12)
        continue

    # ── Layer headings (**Layer N — ...** as standalone line) ─────────────────
    if re.match(r'^\*\*Layer \d', s):
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=6, space_after=2,
                 line_spacing=WD_LINE_SPACING.DOUBLE)
        render_inline(p, s, base_size=12)
        continue

    # ── Multi-line block (normal paragraphs, possibly spanning lines) ─────────
    # Join lines within a block into one paragraph
    lines_in_block = s.splitlines()
    # If it looks like multiple distinct paragraphs within the block, split
    paragraphs_to_add = []
    current_para = []
    for l in lines_in_block:
        ls = l.strip()
        if not ls:
            if current_para:
                paragraphs_to_add.append(' '.join(current_para))
                current_para = []
        else:
            current_para.append(ls)
    if current_para:
        paragraphs_to_add.append(' '.join(current_para))

    for para_text in paragraphs_to_add:
        if not para_text.strip():
            continue
        # Skip metadata lines
        if para_text.startswith('*Written as') or para_text.startswith('### Core'):
            continue
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=0, space_after=0,
                 line_spacing=WD_LINE_SPACING.DOUBLE)
        render_inline(p, para_text, base_size=12)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────

out = '/Users/piyushsingh/Downloads/HECA PROJECT /HECA_Research_Paper_Singh_Chatterjee.docx'
doc.save(out)
print(f'✓  Saved → {out}')
