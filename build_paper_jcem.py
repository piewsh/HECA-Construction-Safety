"""
HECA Research Paper — ASCE JCEM Compliant Word Document
Journal of Construction Engineering and Management

ASCE JCEM Requirements:
- Times New Roman, double-spaced
- Continuous line numbering throughout
- NO numbered sections (word headings only)
- Figures and tables at END of document (not inline)
- Single column
- Abstract ~250 words
- All required sections: Abstract, Introduction, Background,
  Methodology, Results, Discussion, Conclusion, Data Availability,
  Acknowledgements, References
- Author contributions (CRediT)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ─── Colours ──────────────────────────────────────────────────────────────────
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x55, 0x55, 0x55)

# ─── Helpers ──────────────────────────────────────────────────────────────────

def add_line_numbers(doc):
    """Add continuous line numbering to all sections (ASCE requirement)."""
    for section in doc.sections:
        sectPr = section._sectPr
        lnNumType = OxmlElement('w:lnNumType')
        lnNumType.set(qn('w:countBy'), '1')
        lnNumType.set(qn('w:restart'), 'continuous')
        lnNumType.set(qn('w:start'),   '1')
        sectPr.append(lnNumType)

def set_font(run, size=12, bold=False, italic=False):
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK

def body_fmt(p, space_after=0):
    """Standard ASCE body paragraph format."""
    f = p.paragraph_format
    f.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    f.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    f.space_before      = Pt(0)
    f.space_after       = Pt(space_after)
    f.first_line_indent = Inches(0.5)

def heading_fmt(p, space_before=12, space_after=0):
    """ASCE heading: word-only, no numbering, centred or left as appropriate."""
    f = p.paragraph_format
    f.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    f.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    f.space_before      = Pt(space_before)
    f.space_after       = Pt(space_after)
    f.first_line_indent = Pt(0)

def subheading_fmt(p):
    f = p.paragraph_format
    f.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    f.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    f.space_before      = Pt(0)
    f.space_after       = Pt(0)
    f.first_line_indent = Inches(0.5)

def render_inline(para, text, size=12, bold=False, italic=False):
    """Parse **bold**, *italic* and add runs."""
    pat = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))', re.DOTALL)
    for m in pat.finditer(text):
        if m.group(2):
            r = para.add_run(m.group(2))
            set_font(r, size=size, bold=True, italic=italic)
        elif m.group(3):
            r = para.add_run(m.group(3))
            set_font(r, size=size, bold=bold, italic=True)
        elif m.group(4):
            r = para.add_run(m.group(4))
            set_font(r, size=size, bold=bold, italic=italic)

def add_heading(doc, text, level=1):
    """level 1 = major section, level 2 = subsection (italic, left-aligned)."""
    p = doc.add_paragraph()
    if level == 1:
        heading_fmt(p, space_before=24, space_after=0)
        r = p.add_run(text.upper())
        set_font(r, size=12, bold=True)
    else:
        # ASCE subsections: italic, flush left, run-in style
        subheading_fmt(p)
        r = p.add_run(text + '. ')
        set_font(r, size=12, italic=True, bold=False)
    return p

def add_body(doc, text):
    """Add a justified body paragraph."""
    if not text.strip():
        return
    p = doc.add_paragraph()
    body_fmt(p)
    render_inline(p, text)
    return p

def add_blank(doc):
    p = doc.add_paragraph()
    body_fmt(p)
    p.add_run('')

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(doc, caption, rows):
    """Add a table with caption above (ASCE style)."""
    # Caption
    p = doc.add_paragraph()
    body_fmt(p, space_after=2)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(caption)
    set_font(r, size=12, bold=True)

    if not rows:
        return

    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'

    for ri, row_data in enumerate(rows):
        for ci, txt in enumerate(row_data):
            cell = t.cell(ri, ci)
            cell.text = ''
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(3)
            cp.paragraph_format.space_after  = Pt(3)
            cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = cp.add_run(re.sub(r'\*+', '', txt).strip())
            set_font(r, size=10, bold=(ri == 0))
            shade_cell(cell, 'E0E0E0' if ri == 0 else ('F8F8F8' if ri % 2 else 'FFFFFF'))

    p2 = doc.add_paragraph()
    body_fmt(p2)


# ─── Document Setup ───────────────────────────────────────────────────────────

doc = Document()

for s in doc.sections:
    s.page_width    = Inches(8.5)
    s.page_height   = Inches(11)
    s.left_margin   = Inches(1.0)
    s.right_margin  = Inches(1.0)
    s.top_margin    = Inches(1.0)
    s.bottom_margin = Inches(1.0)

# Global normal style
ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(12)
ns.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# ─── TITLE PAGE ───────────────────────────────────────────────────────────────

for _ in range(2):
    add_blank(doc)

# Title
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
r = p.add_run(
    'Measuring What Kills: Energy-Based Safety Assessment of Indian '
    'PPVC Construction Using the HECA Framework'
)
set_font(r, size=14, bold=True)

add_blank(doc)

# Authors (ASCE blind review note)
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
r = p.add_run('Piyush Singh, S.M.ASCE\u00b9  and  Abhishek Chatterjee\u00b9')
set_font(r, size=12)

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
r = p.add_run(
    '\u00b9 Department of Civil Engineering, Indian Institute of Technology Madras, '
    'Chennai 600036, India.'
)
set_font(r, size=12)

add_blank(doc)

# Corresponding author note
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
r = p.add_run('Corresponding Author: piyushsingh@iitm.ac.in')
set_font(r, size=12)

doc.add_page_break()

# ─── ABSTRACT ─────────────────────────────────────────────────────────────────

add_heading(doc, 'Abstract', level=1)

abstract = (
    "India's construction sector accounts for approximately one in every four occupational "
    "fatalities in the country, yet no energy-based safety baseline exists for its rapidly "
    "expanding Prefabricated Prefinished Volumetric Construction (PPVC) sector. This study "
    "deploys the High-Energy Control Assessment (HECA) framework, developed in North America "
    "and never previously applied in India, to an active PPVC project at the Indian Institute "
    "of Technology Madras. Through a non-invasive photographic audit of 320 curated images, "
    "two trained assessors identified 1,583 hazard observations and classified 1,254 (79.2%) "
    "as high-energy with serious injury and fatality (SIF) potential. The aggregate HECA score "
    "was 1.83%, compared to the 40-70% benchmark range reported for North American utility "
    "projects, representing a 20-40 times gap in Direct Control deployment. In 297 of 320 "
    "images (92.8%), every visible high-energy hazard was completely uncontrolled. The PPVC "
    "Module Erection phase recorded 0.00% across 189 high-energy hazard observations; five "
    "of seven energy types observed recorded exactly zero throughout the project. Beyond the "
    "field analysis, the study delivers two technical contributions: a six-layer Image "
    "Anonymization Pipeline achieving compliance with India's Digital Personal Data Protection "
    "(DPDP) Act 2023 for construction site photographs, and a zero-dependency HECA Analytics "
    "Dashboard enabling site-level data exploration without institutional IT infrastructure. "
    "Together, these constitute the first complete, replicable HECA research system for Indian "
    "construction and establish the first energy-disaggregated safety baseline for PPVC "
    "construction anywhere in the world."
)

add_body(doc, abstract)
add_blank(doc)

# Keywords
p = doc.add_paragraph()
body_fmt(p)
p.paragraph_format.first_line_indent = Inches(0.5)
r = p.add_run('Keywords: ')
set_font(r, bold=True)
r = p.add_run(
    'High-Energy Control Assessment (HECA); PPVC; modular construction; '
    'direct controls; SIF prevention; photographic audit; image anonymization; '
    'DPDP Act 2023; India; construction safety'
)
set_font(r, italic=True)

doc.add_page_break()

# ─── BODY CONTENT: parse from markdown ───────────────────────────────────────

with open(
    '/Users/piyushsingh/.gemini/antigravity/brain/'
    '50562d44-3450-4c11-945c-c1fe6417aacf/RESEARCH_PAPER_HUMAN_DRAFT.md',
    encoding='utf-8'
) as f:
    raw = f.read()

# Extract from Section 1 onward
raw = re.sub(r'^.*?(## 1\.)', r'\1', raw, flags=re.DOTALL)
# Drop footer note and references (we handle refs manually)
raw = re.sub(r'\*\[Complete paper.*', '', raw, flags=re.DOTALL)

# Section heading map (ASCE: NO numbers, word headings only)
HEADING_MAP = {
    '1. INTRODUCTION':                             'Introduction',
    '2. BACKGROUND: THE HECA FRAMEWORK IN DEPTH':  'Background',
    '3. METHODOLOGY':                              'Methodology',
    '4. TECHNICAL SYSTEMS: PIPELINE AND DASHBOARD':'Technical Systems: Anonymization Pipeline and Analytics Dashboard',
    '5. RESULTS':                                  'Results',
    '6. DISCUSSION':                               'Discussion',
    '7. CONCLUSION':                               'Conclusion',
    'REFERENCES':                                  'References',
}

# Subsection heading cleanup (drop numbering for sub-levels - ASCE keeps sub numbers)
def clean_subheading(text):
    # ASCE allows sub-numbering like "Research Design" or "1.1 Research Design"
    # We keep the text as-is but remove markdown artefacts
    return text.strip()

# Tables data (extracted already, will be appended at end)
# For ASCE, tables go at the END after references, one per page

TABLE_1 = {
    'caption': 'Table 1. Aggregate HECA Results',
    'rows': [
        ['Metric', 'Value'],
        ['Images analyzed', '320'],
        ['Total hazard observations', '1,583'],
        ['High-energy (HE) hazards', '1,254 (79.2%)'],
        ['Successes (HE with Direct Control)', '23'],
        ['Exposures (HE without Direct Control)', '1,231'],
        ['Aggregate HECA Score', '1.83%'],
        ['Mean HE hazards per image', '3.92 (SD = 0.87)'],
        ['Images with HECA Score = 0%', '297 (92.8%)'],
        ['North American benchmark range', '40-70%'],
    ]
}
TABLE_2 = {
    'caption': 'Table 2. High-Energy Hazard Distribution and Energy-Specific HECA Scores',
    'rows': [
        ['Energy Type', 'Count', '% of HE Total', 'Successes', 'Exposures', 'HECA Score'],
        ['Gravity',     '803',   '64.0%',          '18',        '785',       '2.24%'],
        ['Mechanical',  '135',   '10.8%',          '5',         '130',       '3.70%'],
        ['Electrical',  '122',   '9.7%',           '0',         '122',       '0.00%'],
        ['Motion',      '111',   '8.9%',           '0',         '111',       '0.00%'],
        ['Pressure',    '72',    '5.7%',            '0',         '72',        '0.00%'],
        ['Temperature', '10',    '0.8%',            '0',         '10',        '0.00%'],
        ['Chemical',    '1',     '0.1%',            '0',         '1',         '0.00%'],
        ['Total',       '1,254', '',                '23',        '1,231',     '1.83%'],
    ]
}
TABLE_3 = {
    'caption': 'Table 3. Phase-Wise HECA Performance',
    'rows': [
        ['Phase', 'Images', 'HE Hazards', 'Successes', 'HECA Score', 'Avg HE/Image'],
        ['Foundation',     '40',  '132',  '4',  '2.94%', '3.30'],
        ['Ground Floor',   '69',  '296',  '18', '6.08%', '4.29'],
        ['Upper Floors',   '165', '633',  '1',  '0.16%', '3.84'],
        ['PPVC Erection',  '46',  '189',  '0',  '0.00%', '4.11'],
        ['Total',          '320', '1,254','23', '1.83%', '3.92'],
    ]
}
TABLE_4 = {
    'caption': 'Table 4. SIF Category Analysis (n = 1,254 HE Hazards)',
    'rows': [
        ['SIF Category', 'Observations', '% of HE Total', 'HECA Score'],
        ['Fall-related',                '384', '30.7%', 'approx. 0%'],
        ['Struck-by (swinging load)',   '185', '14.8%', '0.00%'],
        ['Impalement (uncapped rebar)', '150', '12.0%', '0.00%'],
        ['Electrocution',               '114', '9.1%',  '0.00%'],
        ['Crush / pinch',               '98',  '7.8%',  'approx. 0%'],
        ['Other / compound',            '323', '25.8%', 'approx. 0%'],
    ]
}

TABLES_FOR_END = [TABLE_1, TABLE_2, TABLE_3, TABLE_4]

# ── Parse and render ──────────────────────────────────────────────────────────

lines = raw.splitlines()
i = 0
in_table = False
table_buffer = []
skip_inline_tables = True   # tables will be placed at end per ASCE style

while i < len(lines):
    line = lines[i]
    s    = line.strip()

    if not s:
        i += 1
        continue

    # Skip metadata
    if s.startswith('# THE HECA') or s.startswith('### Core') or s.startswith('*Written'):
        i += 1
        continue

    # Horizontal rule — just add blank
    if s == '---':
        add_blank(doc)
        i += 1
        continue

    # Major section heading ##
    if re.match(r'^## ', s):
        raw_title = re.sub(r'^## ', '', s).strip()
        # Map to ASCE heading
        word_heading = HEADING_MAP.get(raw_title, raw_title)
        add_heading(doc, word_heading, level=1)
        i += 1
        continue

    # Subsection heading ###
    if re.match(r'^### ', s):
        text = re.sub(r'^### ', '', s).strip()
        # Remove leading number (ASCE prefers no sub-numbers, or keeps them)
        # We keep them for clarity
        add_heading(doc, text, level=2)
        i += 1
        continue

    # Sub-subsection ####
    if re.match(r'^#### ', s):
        text = re.sub(r'^#### ', '', s).strip()
        p = doc.add_paragraph()
        body_fmt(p)
        r = p.add_run(text + '.  ')
        set_font(r, bold=True, italic=True)
        i += 1
        continue

    # Table — skip inline, tables go at end
    if s.startswith('|'):
        while i < len(lines) and lines[i].strip().startswith('|'):
            i += 1
        # Note: "(See Table N)" will be rendered as inline reference
        continue

    # Figure placeholder
    if re.match(r'^\*?\[Figure', s) or re.match(r'^\[Fig', s):
        txt = re.sub(r'[\*\[\]]', '', s).strip()
        p   = doc.add_paragraph()
        body_fmt(p)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'--- {txt} ---')
        r.font.name   = 'Times New Roman'
        r.font.size   = Pt(11)
        r.font.italic = True
        r.font.color.rgb = GREY
        i += 1
        continue

    # Equation
    if 'HECA Score (%)' in s and '=' in s:
        eq = re.sub(r'\*+', '', s).strip()
        p  = doc.add_paragraph()
        p.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.first_line_indent = Pt(0)
        r  = p.add_run(eq)
        set_font(r, size=12, bold=True)
        i += 1
        continue

    # Bullet list
    if re.match(r'^[-*] ', s):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        render_inline(p, s[2:].strip())
        i += 1
        continue

    # Layer headings
    if re.match(r'^\*\*Layer \d', s):
        p = doc.add_paragraph()
        body_fmt(p)
        render_inline(p, s)
        i += 1
        continue

    # Reference entry [N] ...
    if re.match(r'^\[\d+\]', s):
        p = doc.add_paragraph()
        p.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_before      = Pt(0)
        p.paragraph_format.space_after       = Pt(0)
        p.paragraph_format.left_indent       = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        parts = re.split(r'(\*[^*]+\*)', s)
        for part in parts:
            if part.startswith('*') and part.endswith('*'):
                r = p.add_run(part[1:-1])
                set_font(r, italic=True)
            else:
                r = p.add_run(part)
                set_font(r)
        i += 1
        continue

    # Normal paragraph (join any continuation lines)
    para_lines = [s]
    while i + 1 < len(lines):
        nxt = lines[i+1].strip()
        if (not nxt or nxt.startswith('#') or nxt.startswith('|') or
                nxt == '---' or re.match(r'^[-*] ', nxt) or
                re.match(r'^\[\d+\]', nxt)):
            break
        para_lines.append(nxt)
        i += 1
    full = ' '.join(para_lines)
    add_body(doc, full)
    i += 1

# ─── REQUIRED ASCE SECTIONS ───────────────────────────────────────────────────

doc.add_page_break()
add_heading(doc, 'Data Availability Statement', level=1)
add_body(doc,
    'The anonymized image dataset (320 photographs with HECA coding records) and the '
    'HECA Analytics Dashboard source code generated during this study are available from '
    'the corresponding author upon reasonable request and subject to the anonymization '
    'protocols described in this paper. Raw (non-anonymized) site photographs cannot be '
    'shared due to obligations under the Digital Personal Data Protection Act 2023 (India).'
)

add_blank(doc)
add_heading(doc, 'Acknowledgements', level=1)
add_body(doc,
    'The authors thank the construction project team at IIT Madras for permitting access '
    'to the active project site for photographic data collection. This study was conducted '
    'as part of the Undergraduate Project programme in the Department of Civil Engineering, '
    'Indian Institute of Technology Madras. No external funding was received for this research.'
)

add_blank(doc)
add_heading(doc, 'Author Contributions', level=1)
add_body(doc,
    'Piyush Singh: Conceptualization (co-lead), System design and development (lead), '
    'Image anonymization pipeline (lead), HECA Analytics Dashboard (lead), Data curation '
    '(supporting), Writing (co-lead), Review and editing. '
    'Abhishek Chatterjee: Conceptualization (co-lead), Systematic literature review (lead), '
    'Field data collection (lead), HECA coding and analysis (lead), Formal analysis (lead), '
    'Writing (co-lead), Review and editing.'
)

add_blank(doc)
add_heading(doc, 'Conflict of Interest', level=1)
add_body(doc,
    'The authors declare that they have no known competing financial interests or personal '
    'relationships that could have appeared to influence the work reported in this paper.'
)

# ─── TABLES AT END (ASCE requirement) ────────────────────────────────────────

doc.add_page_break()
add_heading(doc, 'Tables', level=1)

for tdata in TABLES_FOR_END:
    add_table(doc, tdata['caption'], tdata['rows'])
    doc.add_page_break()

# ─── FIGURE CAPTIONS PAGE ─────────────────────────────────────────────────────

add_heading(doc, 'Figure Captions', level=1)
captions = [
    ('Fig. 1.',  'HECA conceptual framework: shift from TRIR-based lagging indicators to energy-based direct control measurement.'),
    ('Fig. 2.',  'HECA Energy Wheel taxonomy showing the ten energy categories and representative Direct Controls for each.'),
    ('Fig. 3.',  'Study site: active PPVC construction project at IIT Madras showing four distinct construction phases.'),
    ('Fig. 4.',  'PRISMA 2020 flow diagram for the systematic literature review (25 included studies, 2009-2026).'),
    ('Fig. 5.',  'Image Anonymization Pipeline architecture showing all six processing layers and the NMS deduplication step.'),
    ('Fig. 6.',  'Sample output from the anonymization pipeline: (a) original image; (b) anonymized output with face blur, OCR masking, and logo removal applied.'),
    ('Fig. 7.',  'HECA Dashboard interface showing the hero KPI section and animated counters for aggregate metrics.'),
    ('Fig. 8.',  'HECA Dashboard analytics section showing energy-type distribution chart with Direct Control status stacked bars and phase-wise comparison with 40% benchmark reference line.'),
    ('Fig. 9.',  'High-energy hazard proportion: (a) 79.2% of all 1,583 observations are high-energy; (b) breakdown by energy type showing gravity dominance at 64.0%.'),
    ('Fig. 10.', 'Phase-wise HECA scores (bars) against North American benchmark range (shaded band, 40-70%), illustrating the performance gap across all four construction phases.'),
    ('Fig. 11.', 'Per-image hazard count distribution histogram showing mean of 3.92 high-energy hazards per image (SD = 0.87).'),
    ('Fig. 12.', 'Gravity sub-type decomposition: five sub-categories within the 803 gravity hazard observations, with Fall from Elevation (18.4%) and Other Gravity (22.1%) dominant.'),
]
for label, cap in captions:
    p = doc.add_paragraph()
    body_fmt(p)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(label + '  ')
    set_font(r, bold=True)
    r2 = p.add_run(cap)
    set_font(r2)

# ─── ADD LINE NUMBERS ─────────────────────────────────────────────────────────

add_line_numbers(doc)

# ─── SAVE ─────────────────────────────────────────────────────────────────────

out = '/Users/piyushsingh/Downloads/HECA PROJECT /HECA_Research_Paper_Singh_Chatterjee_JCEM.docx'
doc.save(out)
print(f'Saved: {out}')
